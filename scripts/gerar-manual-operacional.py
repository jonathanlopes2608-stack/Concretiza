"""Gera docs/MANUAL_OPERACIONAL.docx — manual para analistas e coordenadores."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
MANUAL_DIR = ROOT / "docs" / "manual"
OUT = ROOT / "docs" / "MANUAL_OPERACIONAL.docx"


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        set_run_font(run, size=16 if level == 1 else 13, bold=True, color=RGBColor(0x1E, 0x3A, 0x5F))
    return h


def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run)
        p.paragraph_format.space_after = Pt(2)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_run_font(run)
        p.paragraph_format.space_after = Pt(2)


def add_image(doc, filename, caption):
    path = MANUAL_DIR / filename
    if not path.exists():
        add_para(doc, f"[Imagem não encontrada: {filename}]", bold=True)
        return
    doc.add_picture(str(path), width=Cm(15.5))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    set_run_font(run, size=9, color=RGBColor(0x58, 0x59, 0x5B))
    cap.paragraph_format.space_after = Pt(12)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                set_run_font(run, size=9, bold=True)
    for r_i, row in enumerate(rows):
        cells = table.rows[r_i + 1].cells
        for c_i, val in enumerate(row):
            cells[c_i].text = val
            for p in cells[c_i].paragraphs:
                for run in p.runs:
                    set_run_font(run, size=9)
    doc.add_paragraph()


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Manual Operacional — Concretiza")
    set_run_font(r, size=22, bold=True, color=RGBColor(0x1E, 0x3A, 0x5F))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Pipeline, conformidade e produtividade\nPara analistas e coordenadores")
    set_run_font(r, size=12, color=RGBColor(0x58, 0x59, 0x5B))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("Versão 1.1 — agosto/2026 · App: http://localhost:3047")
    set_run_font(r, size=10)

    # 1 Visão geral
    add_heading(doc, "1. Visão geral", 1)
    add_para(
        doc,
        "O Concretiza controla a fila de produção do correspondente/agente Caixa: "
        "onde cada processo está (fase), por que parou (bloqueio), de quem depende "
        "e o que precisa para seguir — com visão de produtividade no Dashboard.",
    )
    add_para(doc, "Duas camadas importantes:", bold=True)
    add_bullets(
        doc,
        [
            "Fase — posição no funil (Análise, Engenharia, Conformidade, Cartório…).",
            "Bloqueio — trava operacional (“Cliente: enviar CNH”). Pode haver vários abertos na mesma fase.",
        ],
    )
    add_para(doc, "Papéis de acesso:", bold=True)
    add_bullets(
        doc,
        [
            "ADMIN — configuração completa (incluindo tipos de dependência e usuários).",
            "COORDENADOR — fila, reatribuição, força avanço com bloqueio, dashboard.",
            "ANALISTA — trabalha propostas, checklist, abre/resolve bloqueios.",
            "VISUALIZAÇÃO — somente leitura.",
        ],
    )
    add_para(doc, "Acesso de teste (ambiente local):", bold=True)
    add_bullets(
        doc,
        [
            "URL: http://localhost:3047/login",
            "E-mail: admin@concretiza.local",
            "Senha: Admin@123",
        ],
    )

    add_heading(doc, "1.1 Tela de login", 2)
    add_numbered(
        doc,
        [
            "Abra o endereço do sistema.",
            "Informe e-mail e senha.",
            "Clique em Entrar. Se o 2FA estiver ativo, informe o código do autenticador.",
        ],
    )
    add_image(doc, "01-login.png", "Figura 1 — Tela de login")

    # 2 Navegação
    add_heading(doc, "2. Navegação (menu horizontal)", 1)
    add_para(
        doc,
        "Após o login, o sistema usa uma barra superior com módulos (não há menu lateral). "
        "O item ativo fica destacado; à direita aparecem o nome do usuário e o botão Sair. "
        "O logo / nome da marca ficam à esquerda (branding da instalação).",
    )
    add_para(doc, "Módulos do menu:", bold=True)
    add_bullets(
        doc,
        [
            "FILA — lista de processos e detalhe (com abas internas; ver seção 3).",
            "DASHBOARD — produtividade (funil, gargalos, analistas).",
            "AGENDA — compromissos e integração Google Calendar.",
            "USUÁRIOS — gestão de usuários e grupos (ADMIN).",
            "CONFIG — tipos de dependência (“de quem depende”).",
            "CONTA — segurança (2FA) e integrações do usuário.",
        ],
    )
    add_para(
        doc,
        "Dica: o sistema lembra a última rota/filtros de cada módulo na sessão. "
        "Ao clicar de novo em FILA (ou outro módulo), você tende a voltar onde parou.",
    )

    # 3 Abas
    add_heading(doc, "3. Sistema de abas (módulo Fila)", 1)
    add_para(
        doc,
        "Dentro do módulo Fila, abaixo do menu horizontal, há uma barra de abas: "
        "a aba fixa “Fila” (sempre a primeira, sem botão fechar) e abas de processos "
        "abertos (número/nome do processo, com X para fechar).",
    )
    add_para(doc, "Vantagens no dia a dia:", bold=True)
    add_bullets(
        doc,
        [
            "Vários processos abertos ao mesmo tempo — compare, copie dados ou "
            "alterne sem perder o que estava fazendo em cada um.",
            "Fila sempre acessível — a aba fixa “Fila” permanece; um clique volta à lista "
            "com filtros, sem “perder” os processos abertos.",
            "Não perde contexto ao alternar — troca rápida entre lista e detalhes "
            "(e entre processos) sem reabrir tudo do zero a cada clique.",
            "Abas só no módulo Fila — em Dashboard, Agenda, Usuários etc. a barra some "
            "para não poluir outras telas; as abas de processo continuam salvas e "
            "reaparecem ao voltar à Fila.",
            "Estado preservado ao voltar — as abas ficam em armazenamento local do "
            "navegador; fechar o detalhe com X remove só aquela aba.",
        ],
    )
    add_para(doc, "Como usar:", bold=True)
    add_numbered(
        doc,
        [
            "Na Fila, clique no número do processo — abre o detalhe e cria uma aba.",
            "Abra outro processo da mesma forma; as abas ficam lado a lado.",
            "Clique numa aba de processo para alternar; clique em “Fila” para a lista.",
            "Use o X na aba para fechar aquele processo (a Fila nunca fecha).",
        ],
    )
    add_image(doc, "11-abas-processos.png", "Figura 2 — Abas: Fila fixa + processos abertos")

    # 4 Mapeamento
    add_heading(doc, "4. Mapa planilha Excel ↔ POP ↔ sistema", 1)
    add_para(
        doc,
        "A planilha CONTROLE ORÇA/ANÁLISE/ENGENHARIA e o POP (Anexo X — Concessão CCA) "
        "são a referência de negócio. No Concretiza, as fases da coluna “FASE” da planilha "
        "viram a fase do processo; o que antes era anotação em “parecer” vira bloqueio tipado.",
    )
    add_table(
        doc,
        ["Planilha (FASE)", "POP (Anexo X)", "Fase no Concretiza", "O que fazer no sistema"],
        [
            ["ANÁLISE", "Entrevistar / Avaliar cliente", "ANALISE", "Atribuir analista; abrir bloqueio se faltar doc"],
            ["RESTRIÇÃO", "—", "RESTRICAO", "Bloqueio Banco/Cliente + fase Restrição"],
            ["ENGENHARIA", "Avaliar imóvel (OS)", "ENGENHARIA", "Bloqueio Engenharia até OS/laudo"],
            ["DÉBITO FGTS", "Cadastro / FGTS", "DEBITO_FGTS", "Bloqueio Cliente/Banco"],
            ["CONFORMIDADE", "Conformidade da proposta", "CONFORMIDADE", "Checklist documental + bloqueios"],
            ["DECISÃO", "—", "DECISAO", "Aguarda escolha do cliente (banco/condições)"],
            ["EM CARTÓRIO", "Conformidade do registro", "EM_CARTORIO", "Bloqueio Cartório"],
            ["—", "Formalizar contratação", "FORMALIZACAO", "Assinatura / formalização"],
            ["FINALIZADO / CANCELADO", "—", "FINALIZADO / CANCELADO", "Encerrar (motivo se cancelar)"],
        ],
    )
    add_para(
        doc,
        "Dica: tipos novos de “de quem depende” (ex.: Corretor) se cadastram em CONFIG → Dependências — "
        "não é necessário criar uma fase nova para cada ator.",
    )

    # 5 Fila
    add_heading(doc, "5. Fila de processos", 1)
    add_para(
        doc,
        "Objetivo: ver, em uma tela, fase, se está travado, de quem depende, analista, despachante e SLA.",
    )
    add_para(doc, "Como usar:", bold=True)
    add_numbered(
        doc,
        [
            "No menu horizontal, clique em FILA.",
            "Use Busca (processo, Caixa, nome, CPF) e filtros de Fase, Depende de e Analista.",
            "Marque “Só travados” ou “SLA estourado” para priorizar.",
            "Clique no número do processo para abrir o detalhe (abre uma aba).",
            "Use Exportar CSV para baixar a fila atual; Nova proposta / Importar Excel para entrada.",
        ],
    )
    add_para(
        doc,
        "A coluna “Parado em” mostra o bloqueio aberto mais antigo (ex.: Cliente / comprador: Enviar CNH).",
    )
    add_image(doc, "02-fila.png", "Figura 3 — Fila com menu horizontal e barra de abas")

    # 6 Nova proposta
    add_heading(doc, "6. Nova proposta", 1)
    add_para(doc, "Objetivo: cadastrar um processo (equivalente a uma linha da planilha ORÇA).", bold=False)
    add_numbered(
        doc,
        [
            "Em Fila, clique em Nova proposta (ou Importar Excel para lote).",
            "Informe Nº processo interno (ex.: AC/QA 076) e/ou Nº proposta Caixa.",
            "Preencha despachante, modalidade, prioridade e dados do comprador (obrigatórios: nome e CPF).",
            "Clique em Cadastrar proposta — a fase inicial é ENTRADA.",
        ],
    )
    add_image(doc, "03-nova-proposta.png", "Figura 4 — Cadastro manual de proposta")

    # 7 Pipeline
    add_heading(doc, "7. Detalhe — Pipeline (mudar fase)", 1)
    add_para(
        doc,
        "Objetivo: avançar o processo conforme o andamento real (mesmas fases da planilha).",
    )
    add_numbered(
        doc,
        [
            "Abra o processo na fila (ou pela aba já aberta).",
            "No painel Pipeline, escolha a Nova fase e, se for cancelar/reprovar, informe o Motivo.",
            "Clique em Alterar fase.",
            "Se houver bloqueio aberto, o analista deve resolvê-lo; o coordenador/admin pode marcar "
            "“Forçar avanço com bloqueio aberto” (observação obrigatória).",
            "Use Atribuir para definir o analista responsável.",
            "O botão “Linha do tempo” (ao lado do título) mostra o histórico visual do processo.",
        ],
    )
    add_para(
        doc,
        "Regra: para finalizar, não pode haver bloqueio aberto. Ao sair de Conformidade para etapas "
        "seguintes, o checklist documental precisa estar 100% OK.",
    )
    add_image(doc, "04-detalhe-pipeline-bloqueios.png", "Figura 5 — Detalhe do processo (Pipeline + abas)")

    # 8 Bloqueios
    add_heading(doc, "8. Detalhe — Bloqueios (de quem depende)", 1)
    add_para(
        doc,
        "Objetivo: registrar o que falta e de quem depende — substitui anotações soltas na planilha.",
    )
    add_numbered(
        doc,
        [
            "Em Bloqueios / de quem depende, escolha o tipo (Cliente, Despachante, Engenharia…).",
            "Descreva o que precisa para seguir (ex.: Enviar CNH atualizada).",
            "Clique em Registrar / Abrir bloqueio — a fila passa a mostrar isso em “Parado em”.",
            "Quando a pendência for cumprida, clique em Resolver.",
        ],
    )
    add_image(doc, "05-bloqueios.png", "Figura 6 — Bloqueio aberto e formulário para registrar novo")

    # 9 Checklist
    add_heading(doc, "9. Checklist documental (Conformidade)", 1)
    add_para(
        doc,
        "Objetivo: validar documentos do comprador, vendedor e imóvel. Corresponde à etapa de "
        "conformidade do POP (CEOPF/CEOPE) no dia a dia da equipe.",
    )
    add_numbered(
        doc,
        [
            "Na fase Conformidade (ou ao expandir o checklist), anexe o arquivo no item correspondente.",
            "Em RG/CNH, o sistema tenta ler validade e CPF (OCR). Confira o resultado.",
            "Marque OK, PENDENTE ou REPROVADO.",
            "Item REPROVADO abre automaticamente um bloqueio para Cliente.",
            "Item OK fica bloqueado para o analista; só coordenador/admin reabre.",
        ],
    )
    add_image(doc, "06-checklist.png", "Figura 7 — Checklist documental com validação automática")

    # 10 Dashboard
    add_heading(doc, "10. Dashboard de produtividade", 1)
    add_para(doc, "Objetivo: coordenação — funil, gargalos e produção por analista.")
    add_bullets(
        doc,
        [
            "Funil por fase — volume em cada etapa.",
            "Travados por dependência — quem mais segura a fila.",
            "Aging médio — dias na fase atual.",
            "Por analista — ativas, finalizadas no período, SLA estourado.",
            "Filtro 7d / 30d / 90d no topo.",
        ],
    )
    add_para(
        doc,
        "No menu horizontal, clique em DASHBOARD. Nesta tela a barra de abas da Fila não aparece "
        "(as abas de processo ficam guardadas e voltam ao reabrir FILA).",
    )
    add_image(doc, "07-dashboard.png", "Figura 8 — Dashboard de produtividade")

    # 11 Dependencias
    add_heading(doc, "11. Tipos de dependência (CONFIG)", 1)
    add_para(
        doc,
        "Objetivo: manter a lista de “de quem depende” (Cliente, Cartório, Engenharia…). "
        "É possível criar novos tipos sem alterar as fases.",
    )
    add_numbered(
        doc,
        [
            "No menu horizontal, clique em CONFIG (Dependências).",
            "Para criar: informe Código e Nome exibido → Adicionar.",
            "Para ocultar um tipo: Desativar.",
        ],
    )
    add_image(doc, "08-dependencias.png", "Figura 9 — Cadastro de tipos de dependência")

    # 12 Usuarios / Agenda (resumo)
    add_heading(doc, "12. Usuários e Agenda (resumo)", 1)
    add_para(
        doc,
        "USUÁRIOS (ADMIN): cadastrar analistas/coordenadores, definir grupos e permissões de tela/ação.",
    )
    add_image(doc, "09-usuarios.png", "Figura 10 — Gestão de usuários")
    add_para(
        doc,
        "AGENDA: compromissos internos e, se conectado, Google Calendar (conta Gmail pessoal no MVP). "
        "Use CONTA → integrações para vincular o Google.",
    )
    add_image(doc, "10-agenda.png", "Figura 11 — Agenda")

    # 13 Roteiro
    add_heading(doc, "13. Roteiro do dia a dia (da entrada ao finalizado)", 1)
    add_para(doc, "Espelho prático da planilha ORÇA no Concretiza:")
    add_numbered(
        doc,
        [
            "Cadastrar o processo (Nº AC/QA + cliente + despachante).",
            "Atribuir / assumir o analista (fase tende a Análise).",
            "Avançar fases conforme o andamento (Engenharia, Débito FGTS, Conformidade…).",
            "Sempre que depender de terceiro, abrir bloqueio com tipo + o que falta.",
            "Na Conformidade, trabalhar o checklist e anexos.",
            "Resolver bloqueios antes de finalizar.",
            "Use as abas para manter 2–3 processos em paralelo sem perder a Fila.",
            "Acompanhar gargalos no Dashboard; exportar CSV da fila se precisar reportar.",
        ],
    )

    # 14 Glossario
    add_heading(doc, "14. Glossário rápido", 1)
    add_bullets(
        doc,
        [
            "Fase — etapa do pipeline (onde o processo está).",
            "Bloqueio — motivo da parada + de quem depende + o que precisa.",
            "SLA — prazo-alvo da fase atual (destaque vermelho se estourado).",
            "Aba Fila — aba fixa do módulo Fila; sempre disponível enquanto você está em Fila/processo.",
            "Despachante — responsável externo de captação (texto, sem login).",
            "EN QA — analista interno de qualidade/análise.",
            "OS Engenharia — ordem de serviço / vistoria Caixa.",
            "POP — Procedimento Operacional Padrão da Caixa (Anexo X).",
            "Isolve / SIOPI / Caixa Aqui — sistemas externos (contexto); não há integração direta neste MVP.",
        ],
    )

    add_heading(doc, "15. Dicas e erros comuns", 1)
    add_bullets(
        doc,
        [
            "Não avança a fase: verifique bloqueios abertos ou permissão do seu perfil.",
            "Não finaliza: ainda há bloqueio aberto — resolva antes.",
            "Checklist incompleto: não sai de Conformidade para Decisão/Formalização/Cartório.",
            "“Parado em” vazio: não há bloqueio aberto (o processo pode só estar aguardando trabalho interno).",
            "Sumiram as abas? Elas só aparecem em FILA / detalhe de processo — volte pelo menu FILA.",
            "Este Word é editável — atualize prints e textos conforme a operação evoluir.",
            "Regenerar este arquivo: python scripts/gerar-manual-operacional.py "
            "(prints: node scripts/capturar-prints-manual.mjs com o app em 3047).",
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Gerado: {OUT}")


if __name__ == "__main__":
    main()
